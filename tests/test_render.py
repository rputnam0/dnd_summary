from __future__ import annotations

from pathlib import Path

from docx import Document

from dnd_summary.render import render_summary_docx


def test_render_summary_docx_creates_document(tmp_path: Path):
    output = tmp_path / "summary.docx"

    render_summary_docx("First paragraph.\n\nSecond paragraph.", output)

    assert output.exists()
    doc = Document(str(output))
    texts = [para.text for para in doc.paragraphs if para.text]
    assert texts == ["First paragraph.", "Second paragraph."]


def test_render_summary_docx_supports_headings_and_bullets(tmp_path: Path):
    output = tmp_path / "summary.docx"
    content = "# Overview\n\n## Highlights\n- First\n- Second\n\nClosing line."

    render_summary_docx(content, output, title="Session Summary")

    doc = Document(str(output))
    styles = [(para.text, para.style.name) for para in doc.paragraphs if para.text]
    assert styles[0][0] == "Session Summary"
    assert styles[0][1].startswith("Title") or styles[0][1].startswith("Heading")
    assert styles[1][0] == "Overview"
    assert styles[1][1] == "Heading 1"
    assert styles[2][0] == "Highlights"
    assert styles[2][1] == "Heading 2"
    assert styles[3][0] == "First"
    assert styles[4][0] == "Second"
    assert styles[-1][0] == "Closing line."
