from __future__ import annotations

from pathlib import Path

from docx import Document


def _flush_paragraph(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    cleaned = " ".join(line.strip() for line in buffer if line.strip())
    if not cleaned:
        buffer.clear()
        return
    doc.add_paragraph(cleaned)
    buffer.clear()


def render_summary_docx(text: str, output_path: Path, title: str | None = None) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    buffer: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            _flush_paragraph(doc, buffer)
            continue
        if line.startswith("## "):
            _flush_paragraph(doc, buffer)
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            _flush_paragraph(doc, buffer)
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("- "):
            _flush_paragraph(doc, buffer)
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        buffer.append(line)
    _flush_paragraph(doc, buffer)
    doc.save(str(output_path))
